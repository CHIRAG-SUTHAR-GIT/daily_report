from datetime import date
import io
import zipfile

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from src.gujarat_account_formatter import (
    OUTPUT_COLUMNS,
    bankwise_zip_bytes,
    build_kyc_notice_docx,
    build_kyc_notice_pdf,
    dataframe_to_styled_excel_bytes,
    detect_columns,
    prepare_kyc_notice_accounts,
    process_gujarat_account_file,
)


def test_detect_columns_for_automated_workflow_headers():
    columns = ["S No.", "Acknowledgement No.", "IFSC Code", "Bank/FIs", "Account No."]

    detected = detect_columns(columns)

    assert detected == {
        "ACK NO": "Acknowledgement No.",
        "IFSC CODE": "IFSC Code",
        "BANK NAME": "Bank/FIs",
        "AC NO": "Account No.",
    }


def test_process_gujarat_file_removes_duplicate_accounts_and_keeps_blank_suspect_fields():
    df = pd.DataFrame(
        {
            "Acknowledgement No.": ["3111", "3112", "3113", "3114"],
            "IFSC Code": ["SBIN0001", "HDFC0001", "BARB0001", "ICIC0001"],
            "Bank/FIs": ["SBI", "HDFC", "BOB", "ICICI"],
            "Account No.": ["1001", "1001", "", "2002.0"],
        }
    )
    mapping = {
        "ACK NO": "Acknowledgement No.",
        "IFSC CODE": "IFSC Code",
        "BANK NAME": "Bank/FIs",
        "AC NO": "Account No.",
    }

    output, stats = process_gujarat_account_file(df, mapping)

    assert list(output.columns) == OUTPUT_COLUMNS
    assert output["AC NO"].tolist() == ["2002", "1001"]
    assert output["SR NO"].tolist() == [1, 2]
    assert output[OUTPUT_COLUMNS[5:]].eq("").all().all()
    assert stats["duplicate_account_rows"] == 1
    assert stats["blank_account_rows"] == 1
    assert stats["output_rows"] == 2


def test_styled_excel_has_expected_structure():
    df = pd.DataFrame(
        [[1, "3111", "SBIN0001", "SBI", "1001", "", "", "", ""]],
        columns=OUTPUT_COLUMNS,
    )

    excel_bytes = dataframe_to_styled_excel_bytes(df)
    workbook = load_workbook(io.BytesIO(excel_bytes))
    worksheet = workbook.active

    assert [cell.value for cell in worksheet[1]] == OUTPUT_COLUMNS
    assert worksheet.freeze_panes == "A2"
    assert worksheet.auto_filter.ref == "A1:I2"


def test_bankwise_zip_contains_one_excel_per_bank():
    df = pd.DataFrame(
        [
            [1, "3111", "SBIN0001", "SBI", "1001", "", "", "", ""],
            [2, "3112", "HDFC0001", "HDFC", "2002", "", "", "", ""],
        ],
        columns=OUTPUT_COLUMNS,
    )

    zip_bytes = bankwise_zip_bytes(df)

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as archive:
        names = sorted(archive.namelist())
        assert names == ["HDFC.xlsx", "SBI.xlsx"]
        sbi_df = pd.read_excel(io.BytesIO(archive.read("SBI.xlsx")), dtype=str, keep_default_na=False)

    assert list(sbi_df.columns) == OUTPUT_COLUMNS
    assert sbi_df.loc[0, "SR NO"] == "1"
    assert sbi_df.loc[0, "BANK NAME"] == "SBI"


def test_prepare_kyc_notice_accounts_requires_bank_and_ascii_digits():
    df = pd.DataFrame(
        {
            "ACK NO": ["3111", "3112", "3113", "3114", "3115"],
            "BANK NAME": [
                "Example Bank",
                "Example Bank",
                "Example Wallet",
                "Example Bank",
                "EXAMPLE BANK",
            ],
            "AC NO": ["1001.0", "12A45", "2002", "१२३४५", "1001"],
        }
    )

    notice_df, stats = prepare_kyc_notice_accounts(df)

    assert notice_df.to_dict("records") == [
        {
            "SR NO": 1,
            "ACK NO": "3111",
            "AC NO": "1001",
            "BANK NAME": "Example Bank",
        }
    ]
    assert stats == {
        "input_rows": 5,
        "invalid_account_rows": 2,
        "non_bank_rows": 1,
        "duplicate_account_rows": 1,
        "output_rows": 1,
    }


def _kyc_notice_rows(count: int = 55) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "SR NO": index,
                "ACK NO": f"3110926020{index:04d}",
                "AC NO": f"7000000000{index:04d}",
                "BANK NAME": (
                    "Union Bank of India (including Andhra Bank and Corporation Bank)"
                    if index > count - 6
                    else "Example Bank"
                ),
            }
            for index in range(1, count + 1)
        ]
    )


def test_build_kyc_notice_docx_updates_date_rows_and_final_keep_chain():
    notice_df = _kyc_notice_rows()

    docx_bytes = build_kyc_notice_docx(notice_df, date(2026, 9, 3))
    document = Document(io.BytesIO(docx_bytes))
    date_line = next(
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.startswith("No. CCoE/1930/DR/")
    )

    assert "No. CCoE/1930/DR/03/09/2026" in date_line
    assert "Date: 03-09-2026" in date_line
    assert len(document.tables[0].rows) == 56
    assert document.tables[0].rows[1].cells[2].text == "70000000000001"
    assert document.tables[0].rows[-1].cells[2].text == "70000000000055"
    assert all(
        cell.paragraphs[0].paragraph_format.keep_with_next
        for cell in document.tables[0].rows[-1].cells
    )
    closing = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("You are instructed to comply")
    )
    assert closing.paragraph_format.keep_with_next


def test_build_kyc_notice_pdf_keeps_last_account_with_officer_signature_page():
    notice_df = _kyc_notice_rows()

    pdf_bytes = build_kyc_notice_pdf(notice_df, date(2026, 9, 3))
    reader = PdfReader(io.BytesIO(pdf_bytes))
    first_page_text = reader.pages[0].extract_text()
    last_page_text = reader.pages[-1].extract_text()

    assert "No. CCoE/1930/DR/03/09/2026" in first_page_text
    assert "Date: 03-09-2026" in first_page_text
    assert "70000000000055" in last_page_text
    assert "Sanjaykumar Keshvala" in last_page_text
    assert len(reader.pages) >= 2
