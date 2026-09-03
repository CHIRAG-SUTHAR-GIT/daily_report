from copy import deepcopy
from datetime import date, datetime, timedelta, timezone
import io
from pathlib import Path
import re
import zipfile
from typing import Dict, List, Tuple
from xml.sax.saxutils import escape

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Image as PdfImage,
    KeepTogether,
    LongTable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table as PdfTable,
    TableStyle as PdfTableStyle,
)


OUTPUT_COLUMNS = [
    "SR NO",
    "ACK NO",
    "IFSC CODE",
    "BANK NAME",
    "AC NO",
    "ACCOUNT HOLDER'S NAME",
    "ACCOUNT HOLDER'S ADDRESS",
    "ACCOUNT HOLDER'S MOBILE NUMBER",
    "ACCOUNT HOLDER'S LOCATION",
]

SOURCE_FIELDS = {
    "ACK NO": [
        "acknowledgementno",
        "acknowledgementnumber",
        "acknowledgmentno",
        "ackno",
        "acknumber",
        "ack",
    ],
    "IFSC CODE": [
        "ifsccode",
        "ifsc",
    ],
    "BANK NAME": [
        "bankfis",
        "bankfi",
        "bankname",
        "bank",
        "bankfinancialinstitution",
    ],
    "AC NO": [
        "accountno",
        "accountnumber",
        "bankaccountno",
        "bankaccountnumber",
        "acno",
        "accno",
        "account",
    ],
}

_ILLEGAL_EXCEL_CHARS_RE = re.compile(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]")
_INVALID_FILENAME_CHARS_RE = re.compile(r'[<>:"/\\|?*\x00-\x1F]+')
_ASCII_ACCOUNT_RE = re.compile(r"[0-9]+")
_INDIA_TIMEZONE = timezone(timedelta(hours=5, minutes=30))
_KYC_NOTICE_TEMPLATE = (
    Path(__file__).resolve().parent.parent / "assets" / "kyc_notice_template.docx"
)
_KYC_NOTICE_COLUMNS = ["ACK NO", "BANK NAME", "AC NO"]
_KYC_TABLE_WIDTHS = [0.48 * inch, 1.15 * inch, 2.84 * inch, 2.33 * inch]


def _normalize_header(value) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value).lower())


def _clean_text(value) -> str:
    if value is None:
        return ""
    try:
        if value != value:
            return ""
    except Exception:
        pass
    text = str(value).strip()
    if text.lower() in {"nan", "none", "<na>", "nat"}:
        return ""
    return _ILLEGAL_EXCEL_CHARS_RE.sub("", text)


def _clean_identifier(value) -> str:
    text = _clean_text(value)
    if re.fullmatch(r"\d+\.0", text):
        return text[:-2]
    if re.fullmatch(r"\d+(?:\.0+)?[eE]\+?\d+", text):
        try:
            return str(int(float(text)))
        except (OverflowError, ValueError):
            return text
    return text


def _account_key(value) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "", _clean_identifier(value)).upper()


def detect_columns(columns: List[str]) -> Dict[str, str]:
    normalized = {_normalize_header(column): column for column in columns}
    detected: Dict[str, str] = {}
    for field, aliases in SOURCE_FIELDS.items():
        for alias in aliases:
            if alias in normalized:
                detected[field] = normalized[alias]
                break
    return detected


def process_gujarat_account_file(
    df: pd.DataFrame,
    column_mapping: Dict[str, str],
) -> Tuple[pd.DataFrame, Dict[str, int]]:
    missing = [
        field
        for field in ("ACK NO", "IFSC CODE", "BANK NAME", "AC NO")
        if not column_mapping.get(field) or column_mapping[field] not in df.columns
    ]
    if missing:
        raise ValueError(f"Missing source column mapping for: {', '.join(missing)}")

    output = pd.DataFrame(
        {
            "ACK NO": df[column_mapping["ACK NO"]].map(_clean_identifier),
            "IFSC CODE": df[column_mapping["IFSC CODE"]].map(_clean_identifier),
            "BANK NAME": df[column_mapping["BANK NAME"]].map(_clean_text),
            "AC NO": df[column_mapping["AC NO"]].map(_clean_identifier),
            "ACCOUNT HOLDER'S NAME": "",
            "ACCOUNT HOLDER'S ADDRESS": "",
            "ACCOUNT HOLDER'S MOBILE NUMBER": "",
            "ACCOUNT HOLDER'S LOCATION": "",
        }
    )

    original_rows = len(output)
    output["_account_key"] = output["AC NO"].map(_account_key)
    blank_account_rows = int((output["_account_key"] == "").sum())
    output = output[output["_account_key"] != ""].copy()

    before_dedupe = len(output)
    output = output.drop_duplicates(subset="_account_key", keep="first").copy()
    duplicate_account_rows = before_dedupe - len(output)

    output = output.sort_values(
        by=["BANK NAME", "AC NO"],
        key=lambda series: series.astype(str).str.upper(),
        kind="stable",
    ).reset_index(drop=True)

    output.insert(0, "SR NO", range(1, len(output) + 1))
    output = output[OUTPUT_COLUMNS]

    stats = {
        "input_rows": original_rows,
        "blank_account_rows": blank_account_rows,
        "duplicate_account_rows": duplicate_account_rows,
        "output_rows": len(output),
        "bank_count": int(output["BANK NAME"].replace("", "UNKNOWN BANK").nunique()),
    }
    return output, stats


def _safe_sheet_name(value: str) -> str:
    name = _INVALID_FILENAME_CHARS_RE.sub("_", _clean_text(value)).strip(" .'_")
    return (name or "Data")[:31]


def _safe_filename(value: str) -> str:
    name = _INVALID_FILENAME_CHARS_RE.sub("_", _clean_text(value)).strip(" ._")
    return name or "UNKNOWN_BANK"


def _style_worksheet(worksheet, table_name: str) -> None:
    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    body_font = Font(name="Calibri", size=10, color="111827")
    odd_fill = PatternFill("solid", fgColor="FFFFFF")
    even_fill = PatternFill("solid", fgColor="F4F8FB")
    blank_fill = PatternFill("solid", fgColor="FFF7DA")
    border_side = Side(style="thin", color="D9E2F3")
    border = Border(left=border_side, right=border_side, top=border_side, bottom=border_side)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="top", wrap_text=True)

    worksheet.freeze_panes = "A2"
    worksheet.sheet_view.showGridLines = False
    worksheet.row_dimensions[1].height = 28

    for cell in worksheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    for row_index, row in enumerate(worksheet.iter_rows(min_row=2), start=2):
        row_fill = even_fill if row_index % 2 == 0 else odd_fill
        worksheet.row_dimensions[row_index].height = 24
        for column_index, cell in enumerate(row, start=1):
            cell.font = body_font
            cell.border = border
            cell.fill = blank_fill if column_index >= 6 else row_fill
            cell.alignment = center if column_index in {1, 2, 3, 5, 7} else left
            if column_index in {2, 3, 5, 7}:
                cell.number_format = "@"

    preferred_widths = {
        "A": 8,
        "B": 22,
        "C": 16,
        "D": 30,
        "E": 24,
        "F": 24,
        "G": 22,
        "H": 38,
        "I": 28,
    }
    for column_cells in worksheet.columns:
        letter = column_cells[0].column_letter
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        worksheet.column_dimensions[letter].width = min(max(preferred_widths.get(letter, 12), max_length + 3), 48)

    if worksheet.max_row > 1:
        table_ref = f"A1:I{worksheet.max_row}"
        worksheet.auto_filter.ref = table_ref
        table = Table(displayName=table_name, ref=table_ref)
        table.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=False,
            showColumnStripes=False,
        )
        worksheet.add_table(table)
    else:
        worksheet.auto_filter.ref = "A1:I1"


def dataframe_to_styled_excel_bytes(df: pd.DataFrame, sheet_name: str = "Unique Accounts") -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=_safe_sheet_name(sheet_name))
        worksheet = writer.sheets[_safe_sheet_name(sheet_name)]
        _style_worksheet(worksheet, "UniqueAccountsTable")
    return buffer.getvalue()


def bankwise_zip_bytes(df: pd.DataFrame) -> bytes:
    zip_buffer = io.BytesIO()
    used_names = set()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        grouped = df.assign(_bank=df["BANK NAME"].replace("", "UNKNOWN BANK")).groupby("_bank", sort=True)
        for bank_name, bank_df in grouped:
            bank_output = bank_df.drop(columns=["_bank"]).copy().reset_index(drop=True)
            bank_output["SR NO"] = range(1, len(bank_output) + 1)
            excel_bytes = dataframe_to_styled_excel_bytes(bank_output, sheet_name=str(bank_name))

            filename_root = _safe_filename(str(bank_name))
            filename = f"{filename_root}.xlsx"
            suffix = 2
            while filename.lower() in used_names:
                filename = f"{filename_root}_{suffix}.xlsx"
                suffix += 1
            used_names.add(filename.lower())
            archive.writestr(filename, excel_bytes)
    return zip_buffer.getvalue()


def _current_notice_date() -> date:
    return datetime.now(_INDIA_TIMEZONE).date()


def prepare_kyc_notice_accounts(
    df: pd.DataFrame,
) -> Tuple[pd.DataFrame, Dict[str, int]]:
    """Keep unique numeric accounts whose bank name contains ``Bank``."""
    missing = [column for column in _KYC_NOTICE_COLUMNS if column not in df.columns]
    if missing:
        raise ValueError(f"Missing KYC notice column(s): {', '.join(missing)}")

    notice_df = pd.DataFrame(
        {
            "ACK NO": df["ACK NO"].map(_clean_identifier),
            "BANK NAME": df["BANK NAME"].map(_clean_text),
            "AC NO": df["AC NO"].map(_clean_identifier),
        }
    )
    ascii_account_mask = notice_df["AC NO"].map(
        lambda value: _ASCII_ACCOUNT_RE.fullmatch(value) is not None
    )
    bank_name_mask = notice_df["BANK NAME"].str.contains(
        "BANK", case=False, regex=False, na=False
    )

    invalid_account_rows = int((~ascii_account_mask).sum())
    non_bank_rows = int((ascii_account_mask & ~bank_name_mask).sum())
    notice_df = notice_df[ascii_account_mask & bank_name_mask].copy()
    before_dedupe = len(notice_df)
    notice_df = notice_df.drop_duplicates(subset="AC NO", keep="first").copy()
    duplicate_account_rows = before_dedupe - len(notice_df)
    notice_df.insert(0, "SR NO", range(1, len(notice_df) + 1))
    notice_df = notice_df[["SR NO", "ACK NO", "AC NO", "BANK NAME"]].reset_index(
        drop=True
    )

    return notice_df, {
        "input_rows": len(df),
        "invalid_account_rows": invalid_account_rows,
        "non_bank_rows": non_bank_rows,
        "duplicate_account_rows": duplicate_account_rows,
        "output_rows": len(notice_df),
    }


def _set_word_run_font(run, name: str, size: int, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run_properties = run._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:ascii"), name)
    run_properties.rFonts.set(qn("w:hAnsi"), name)


def _set_notice_table_cell(cell, value: object, alignment) -> None:
    cell.text = _clean_text(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.keep_together = True
    for run in paragraph.runs:
        _set_word_run_font(run, "Calibri", 10)


def _set_cant_split(row) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    if table_row_properties.find(qn("w:cantSplit")) is None:
        table_row_properties.append(OxmlElement("w:cantSplit"))


def _replace_notice_date(document: Document, notice_date: date) -> None:
    paragraph = next(
        (
            item
            for item in document.paragraphs
            if item.text.strip().startswith("No. CCoE/1930/DR/")
        ),
        None,
    )
    if paragraph is None:
        raise ValueError("The KYC notice template date line was not found.")

    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.35), WD_TAB_ALIGNMENT.RIGHT
    )
    reference_run = paragraph.add_run(
        f"No. CCoE/1930/DR/{notice_date.strftime('%d/%m/%Y')}"
    )
    separator_run = paragraph.add_run("\t")
    date_run = paragraph.add_run(f"Date: {notice_date.strftime('%d-%m-%Y')}")
    for run in (reference_run, separator_run, date_run):
        _set_word_run_font(run, "Times New Roman", 12, bold=True)


def _replace_notice_table(document: Document, notice_df: pd.DataFrame) -> None:
    if not document.tables or len(document.tables[0].rows) < 2:
        raise ValueError("The KYC notice template account table is missing.")

    table = document.tables[0]
    template_row = deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    for record in notice_df.itertuples(index=False):
        table._tbl.append(deepcopy(template_row))
        row = table.rows[-1]
        _set_cant_split(row)
        _set_notice_table_cell(row.cells[0], f"{record[0]}.", WD_ALIGN_PARAGRAPH.RIGHT)
        _set_notice_table_cell(row.cells[1], record[1], WD_ALIGN_PARAGRAPH.CENTER)
        _set_notice_table_cell(row.cells[2], record[2], WD_ALIGN_PARAGRAPH.CENTER)
        _set_notice_table_cell(row.cells[3], record[3], WD_ALIGN_PARAGRAPH.LEFT)

    final_row = table.rows[-1]
    for cell in final_row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = True

    closing_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip().startswith("You are instructed to comply")
        ),
        None,
    )
    officer_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if "Sanjaykumar Keshvala" in paragraph.text
        ),
        None,
    )
    if closing_index is None or officer_index is None:
        raise ValueError("The KYC notice template signature block was not found.")

    start_index = max(0, closing_index - 1)
    for paragraph in document.paragraphs[start_index:officer_index]:
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = True
    document.paragraphs[officer_index].paragraph_format.keep_together = True


def build_kyc_notice_docx(
    notice_df: pd.DataFrame,
    notice_date: date | None = None,
) -> bytes:
    """Build a Word KYC notice from the retained official notice template."""
    notice_df, _ = prepare_kyc_notice_accounts(notice_df)
    if notice_df.empty:
        raise ValueError("No eligible bank accounts are available for the KYC notice.")
    if not _KYC_NOTICE_TEMPLATE.exists():
        raise FileNotFoundError("The KYC notice Word template is missing.")

    generated_on = notice_date or _current_notice_date()
    document = Document(_KYC_NOTICE_TEMPLATE)
    _replace_notice_date(document, generated_on)
    _replace_notice_table(document, notice_df)
    document.core_properties.created = datetime.combine(
        generated_on, datetime.min.time()
    )
    document.core_properties.modified = datetime.now(_INDIA_TIMEZONE).replace(
        tzinfo=None
    )
    document.core_properties.author = "Cyber Centre of Excellence, Gujarat"
    document.core_properties.last_modified_by = "Cyber Centre of Excellence, Gujarat"
    document.core_properties.title = (
        f"DR KYC {generated_on.strftime('%d%m%Y')} Bank Notice"
    )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _template_image_bytes(filename: str) -> bytes:
    with zipfile.ZipFile(_KYC_NOTICE_TEMPLATE) as archive:
        return archive.read(f"word/media/{filename}")


def _pdf_styles() -> Dict[str, ParagraphStyle]:
    return {
        "office": ParagraphStyle(
            "KycOffice",
            fontName="Times-Roman",
            fontSize=11,
            leading=15,
            alignment=TA_CENTER,
            spaceAfter=0,
        ),
        "notice_title": ParagraphStyle(
            "KycNoticeTitle",
            fontName="Times-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#EE0000"),
            alignment=TA_CENTER,
            spaceBefore=16,
            spaceAfter=20,
        ),
        "body": ParagraphStyle(
            "KycBody",
            fontName="Times-Roman",
            fontSize=11.5,
            leading=15,
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "justified": ParagraphStyle(
            "KycJustified",
            fontName="Times-Roman",
            fontSize=11.5,
            leading=15,
            alignment=TA_JUSTIFY,
            firstLineIndent=18,
            spaceAfter=8,
        ),
        "table_center": ParagraphStyle(
            "KycTableCenter",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            alignment=TA_CENTER,
        ),
        "table_left": ParagraphStyle(
            "KycTableLeft",
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            alignment=TA_LEFT,
        ),
        "officer": ParagraphStyle(
            "KycOfficer",
            fontName="Times-Roman",
            fontSize=11.5,
            leading=14,
            alignment=TA_CENTER,
        ),
    }


def _pdf_account_row(record, styles: Dict[str, ParagraphStyle]) -> List[Paragraph]:
    values = [
        f"{record[0]}.",
        _clean_text(record[1]),
        _clean_text(record[2]),
        _clean_text(record[3]),
    ]
    return [
        Paragraph(escape(values[0]), styles["table_center"]),
        Paragraph(escape(values[1]), styles["table_center"]),
        Paragraph(escape(values[2]), styles["table_center"]),
        Paragraph(escape(values[3]), styles["table_left"]),
    ]


def _pdf_account_table(rows: List[List[Paragraph]], include_header: bool = False):
    table_rows: List[List[object]] = []
    if include_header:
        table_rows.append(
            [
                Paragraph("<b>Sr.<br/>No.</b>", _pdf_styles()["table_center"]),
                Paragraph("<b>ACK. No.</b>", _pdf_styles()["table_center"]),
                Paragraph("<b>Account No.</b>", _pdf_styles()["table_center"]),
                Paragraph("<b>Bank Name</b>", _pdf_styles()["table_center"]),
            ]
        )
    table_rows.extend(rows)
    table = LongTable(
        table_rows,
        colWidths=_KYC_TABLE_WIDTHS,
        repeatRows=0,
        splitByRow=1,
        hAlign="LEFT",
    )
    table.setStyle(
        PdfTableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B7B7B7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
            + (
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9")),
                    ("TOPPADDING", (0, 0), (-1, 0), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 9),
                ]
                if include_header
                else []
            )
        )
    )
    return table


def build_kyc_notice_pdf(
    notice_df: pd.DataFrame,
    notice_date: date | None = None,
) -> bytes:
    """Build a PDF KYC notice with the final account row kept with the signature."""
    notice_df, _ = prepare_kyc_notice_accounts(notice_df)
    if notice_df.empty:
        raise ValueError("No eligible bank accounts are available for the KYC notice.")

    generated_on = notice_date or _current_notice_date()
    styles = _pdf_styles()
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=0.9 * inch,
        rightMargin=0.58 * inch,
        topMargin=0.18 * inch,
        bottomMargin=0.42 * inch,
        title=f"DR KYC {generated_on.strftime('%d%m%Y')} Bank Notice",
        author="Cyber Centre of Excellence, Gujarat",
    )

    logo = PdfImage(io.BytesIO(_template_image_bytes("image1.png")))
    logo.drawHeight = 1.24 * inch
    logo.drawWidth = 0.96 * inch
    office = Paragraph(
        "<b><u>O/o the Inspector General of Police,<br/>"
        "Cyber Centre of Excellence,</u></b><br/>"
        "7th Floor, Block-2, Karmyogi Bhavan,<br/>"
        "Sector-10 A, Gandhinagar-382010,<br/>"
        "Phone No: 079-23250798<br/>"
        "E-Mail: <u>helpline-cyber-cid@gujarat.gov.in</u>",
        styles["office"],
    )
    header = PdfTable([[logo, office]], colWidths=[1.15 * inch, 5.65 * inch])
    header.setStyle(
        PdfTableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    reference_line = PdfTable(
        [
            [
                Paragraph(
                    f"<b>No. CCoE/1930/DR/{generated_on.strftime('%d/%m/%Y')}</b>",
                    styles["body"],
                ),
                Paragraph(
                    f"<b>Date: {generated_on.strftime('%d-%m-%Y')}</b>",
                    ParagraphStyle(
                        "KycReferenceDate",
                        parent=styles["body"],
                        alignment=TA_RIGHT,
                    ),
                ),
            ]
        ],
        colWidths=[4.45 * inch, 2.35 * inch],
    )
    reference_line.setStyle(
        PdfTableStyle(
            [
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )

    records = list(notice_df.itertuples(index=False, name=None))
    pdf_rows = [_pdf_account_row(record, styles) for record in records]
    story = [
        header,
        HRFlowable(width="100%", thickness=1.4, color=colors.black, spaceBefore=3),
        reference_line,
        Paragraph(
            "<u>Notice u/s 94 BNSS read with Section 168 BNSS</u>",
            styles["notice_title"],
        ),
        Paragraph("To,", styles["body"]),
        Paragraph("All Nodal Officer,", styles["body"]),
        Paragraph(
            "&nbsp;&nbsp;&nbsp;&nbsp;<b>Subject:</b> Provision of Know Your Customer "
            "(KYC) Details of Account Holder(s).",
            styles["body"],
        ),
        Paragraph(
            "This office has received an intimation through NCCRP regarding a "
            "suspicious/fraudulent transaction linked to the bank account(s) "
            "mentioned below. In this connection, you are hereby directed to "
            "furnish the Know Your Customer (KYC) details - including Name, "
            "Address, and Mobile Number - of the concerned account holder(s) at "
            "the earliest.",
            styles["justified"],
        ),
        Paragraph("<b>Suspect Account Details:</b>", styles["body"]),
    ]

    if len(pdf_rows) > 1:
        story.append(_pdf_account_table(pdf_rows[:-1], include_header=True))
        final_table = _pdf_account_table([pdf_rows[-1]], include_header=False)
    else:
        story.append(_pdf_account_table([], include_header=True))
        final_table = _pdf_account_table(pdf_rows, include_header=False)

    closing = Paragraph(
        "You are instructed to comply with this notice and provide the above "
        "information without delay, as required under Section 168 BNSS read "
        "with Section 94 BNSS.",
        ParagraphStyle(
            "KycClosing",
            parent=styles["body"],
            alignment=TA_JUSTIFY,
            spaceAfter=0,
        ),
    )
    signature = PdfImage(io.BytesIO(_template_image_bytes("image2.png")))
    signature.drawWidth = 2.58 * inch
    signature.drawHeight = 1.08 * inch
    officer = Paragraph(
        "<b>(Sanjaykumar Keshvala),</b><br/>"
        "Superintendent of Police,<br/>"
        "Cyber Center of Excellence,<br/>"
        "Gandhinagar, Gujarat State.",
        styles["officer"],
    )
    signature_block = PdfTable(
        [[signature], [officer]],
        colWidths=[2.9 * inch],
        hAlign="RIGHT",
    )
    signature_block.setStyle(
        PdfTableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(
        KeepTogether(
            [
                final_table,
                Spacer(1, 18),
                closing,
                Spacer(1, 12),
                signature_block,
            ]
        )
    )
    document.build(story)
    return output.getvalue()


def _read_uploaded_excel(uploaded_file, sheet_name: str) -> pd.DataFrame:
    uploaded_file.seek(0)
    return pd.read_excel(uploaded_file, sheet_name=sheet_name, dtype=str, keep_default_na=False)


def _select_source_column(label: str, columns: List[str], detected: Dict[str, str]) -> str:
    options = ["-- Select --"] + columns
    detected_column = detected.get(label)
    index = options.index(detected_column) if detected_column in options else 0
    return st.selectbox(label, options=options, index=index, key=f"guj_formatter_{label}")


def render_gujarat_account_formatter_page():
    st.title("Gujarat Unique Account Output")

    uploaded_file = st.file_uploader(
        "Upload Gujarat Excel File",
        type=["xlsx", "xls"],
        key="gujarat_account_formatter_upload",
    )

    if uploaded_file is None:
        st.info("Upload the Gujarat Excel generated from Automated Workflow.")
        return

    try:
        file_bytes = uploaded_file.getvalue()
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
        sheet_name = excel_file.sheet_names[0]
        if len(excel_file.sheet_names) > 1:
            sheet_name = st.selectbox(
                "Sheet",
                options=excel_file.sheet_names,
                key="gujarat_account_formatter_sheet",
            )

        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name, dtype=str, keep_default_na=False)
        st.success(f"Loaded {len(df):,} rows from {uploaded_file.name}")

        detected = detect_columns(list(df.columns))
        st.markdown("### Source Columns")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            ack_col = _select_source_column("ACK NO", list(df.columns), detected)
        with col2:
            ifsc_col = _select_source_column("IFSC CODE", list(df.columns), detected)
        with col3:
            bank_col = _select_source_column("BANK NAME", list(df.columns), detected)
        with col4:
            account_col = _select_source_column("AC NO", list(df.columns), detected)

        selected = {
            "ACK NO": ack_col,
            "IFSC CODE": ifsc_col,
            "BANK NAME": bank_col,
            "AC NO": account_col,
        }
        if any(value == "-- Select --" for value in selected.values()):
            st.warning("Select all four source columns to generate output.")
            return

        output_df, stats = process_gujarat_account_file(df, selected)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_data = dataframe_to_styled_excel_bytes(output_df)
        zip_data = bankwise_zip_bytes(output_df) if not output_df.empty else b""
        notice_date = _current_notice_date()
        notice_df, notice_stats = prepare_kyc_notice_accounts(output_df)
        notice_docx = (
            build_kyc_notice_docx(notice_df, notice_date)
            if not notice_df.empty
            else b""
        )
        notice_pdf = (
            build_kyc_notice_pdf(notice_df, notice_date)
            if not notice_df.empty
            else b""
        )
        notice_filename = f"DR KYC {notice_date.strftime('%d%m%Y')} Bank Notice"

        st.markdown("### Output Summary")
        metric_cols = st.columns(5)
        metric_cols[0].metric("Input Rows", f"{stats['input_rows']:,}")
        metric_cols[1].metric("Output Rows", f"{stats['output_rows']:,}")
        metric_cols[2].metric("Duplicate AC Removed", f"{stats['duplicate_account_rows']:,}")
        metric_cols[3].metric("Blank AC Removed", f"{stats['blank_account_rows']:,}")
        metric_cols[4].metric("Banks", f"{stats['bank_count']:,}")

        st.caption(
            f"KYC notice date: {notice_date.strftime('%d-%m-%Y')} (today) | "
            f"Eligible notice accounts: {notice_stats['output_rows']:,} | "
            f"Non-numeric account rows excluded: "
            f"{notice_stats['invalid_account_rows']:,} | "
            f"Bank names without 'Bank' excluded: {notice_stats['non_bank_rows']:,}"
        )

        st.dataframe(output_df.head(100), use_container_width=True, hide_index=True)

        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.download_button(
                "Download Proper Excel",
                data=excel_data,
                file_name=f"Gujarat_Unique_Accounts_{timestamp}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary",
                use_container_width=True,
                key="gujarat_account_formatter_excel",
            )
        with dl_col2:
            st.download_button(
                "Download Bank-Wise ZIP",
                data=zip_data,
                file_name=f"Gujarat_Bank_Wise_Unique_Accounts_{timestamp}.zip",
                mime="application/zip",
                type="primary",
                use_container_width=True,
                disabled=output_df.empty,
                key="gujarat_account_formatter_zip",
            )

        notice_col1, notice_col2 = st.columns(2)
        with notice_col1:
            st.download_button(
                "Download KYC Notice (Word)",
                data=notice_docx,
                file_name=f"{notice_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
                disabled=notice_df.empty,
                key="gujarat_account_formatter_kyc_docx",
            )
        with notice_col2:
            st.download_button(
                "Download KYC Notice (PDF)",
                data=notice_pdf,
                file_name=f"{notice_filename}.pdf",
                mime="application/pdf",
                type="primary",
                use_container_width=True,
                disabled=notice_df.empty,
                key="gujarat_account_formatter_kyc_pdf",
            )
        if notice_df.empty:
            st.warning(
                "No KYC notice was generated because no row had both a bank "
                "name containing 'Bank' and an account number made only of "
                "ASCII digits 0-9."
            )
    except Exception as exc:
        st.error(f"Error processing file: {exc}")
